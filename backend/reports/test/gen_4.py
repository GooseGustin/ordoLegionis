from docx import Document 
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt
from docx.oxml import ns, OxmlElement
from docx.oxml.ns import qn, nsmap
from math import ceil
from report_variables import *

from datetime import date
currency = '\u20A6'

def format_date(date_string): 
    months = [
        'Jan', 'Feb', 'Mar', 
        'Apr', 'May', 'Jun', 
        'Jul', 'Aug', 'Sep', 
        'Oct', 'Nov', 'Dec'
    ]
    if date_string: 
        day_present = True if date_string.count('-')==2 else False
        if not day_present: date_string += '-10'
        date_obj = date(*[int(i) for i in date_string.split('-')])
        # print(date_obj, date_obj.month)
        [month, day, year] = months[date_obj.month-1], date_obj.day, date_obj.year 
        formatted_date = ''
        if day_present:
            formatted_date = f'{day} {month}, {year}'
        else: 
            formatted_date = f'{month} {year}'
        return [formatted_date, month, day, year] 
    return ['', '', '', '']

def get_term(app_date, sub_date): 
    # Takes in the appointment date and submission date in order to calculate the officer trm 
    appDate = date(*[int(i) for i in app_date.split('-')])
    subDate = date(*[int(i) for i in sub_date.split('-')])
    years = subDate.year - appDate.year
    has_anniversary_passed = (subDate.month > appDate.month) or (subDate.month == appDate.month and subDate.month >= appDate.day)
    if has_anniversary_passed:
        years-=1
    term = years//3 + 1
    if term <= 3: 
        return term 
    return 3

def set_table_border(table, color='FFFFFF'): 
    """Apply custom border colour to a table in python-docx"""
    tbl = table._element # Access table XML element
    tblPr = tbl.find(ns.qn("w:tblPr")) # find table properties

    # Create <w:tblBorders> element
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]: 
        border = OxmlElement(f"w:{border_name}")
        border.set(ns.qn("w:val"), 'single') # single line border
        border.set(ns.qn('w:sz'), '12') # width
        border.set(ns.qn('w:color'), color) 
        border.set(ns.qn('w:space'), '0') # no spacing
        tblBorders.append(border) 

    tblPr.append(tblBorders) # Apply borders to table 

def set_cell_content(cell, text, bold=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 0
    run = p.add_run(text)
    run.bold = bold

def generate_report_docx(curia, praesidium, report):
    doc = Document()
    # Format text 
    style = doc.styles['Normal']
    style.font.name = 'Arial' # 'Times New Roman'
    style.font.size = Pt(9)
    # Ensure the font applies correctly by setting it explicitly for each run
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial') # 'Times New Roman')
    
    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    # section.left_margin = Inches(.5) 
    # section.right_margin = Inches(.5) 
    section.bottom_margin = Inches(.5) 

    # Set page numbers
    def create_element(name):
        return OxmlElement(name) 

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = 'PAGE'

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number(paragraph.add_run())
    # add_page_number(section.footer.paragraphs[0].add_run())


    # Create a table with one row and three columns
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False  # Prevent auto-resizing
    table.style = "Table Grid"
    set_table_border(table)

    # Set column widths (adjust as needed)
    table.rows[0].cells[0].width = Inches(1.1)
    table.rows[0].cells[1].width = Inches(5)
    table.rows[0].cells[2].width = Inches(3)
    # table.columns[0].width = Inches(.5)  # Left column (image)
    # table.columns[1].width = Inches(7)    # Middle column (text)
    # table.columns[2].width = Inches(.5)  # Right column (empty or another element)

    # Insert image in the first cell (left side)
    cell_1 = table.cell(0, 0)
    paragraph_1 = cell_1.paragraphs[0]
    run_1 = paragraph_1.add_run()
    run_1.add_picture("./Vexillium_Legionis.png", width=Inches(1), height=Inches(1.5))  # Adjust as needed

    # Insert multi-line heading in the middle cell
    cell_2 = table.cell(0, 1)
    paragraph_2 = cell_2.paragraphs[0]
    run_2 = paragraph_2.add_run("LEGIONIS MARIAE\n")
    run_2.bold = True
    run_2.font.size = Pt(14)
    heading_text = curia['name'].upper() + '\n'
    heading_text += curia['parish'].upper() + '\n\n'
    heading_text += f"EMAIL: {curia['email'].lower()}\n\n"
    heading_text += "PRAESIDIUM ANNUAL REPORT"
    run_2 = paragraph_2.add_run(heading_text)
    run_2.bold = True
    run_2.font.size = Pt(10)  # Adjust font size

    # Center-align text in the middle cell (both vertically and horizontally)
    tc = cell_2._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)
    paragraph_2.alignment = 1  # Center-align text horizontally

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Date of Report: ")
    run.bold = True 
    p.add_run(format_date(report['submission_date'])[0])
    p.add_run("\t\t\t\t\tDate of Last Report: ").bold = True 
    p.add_run(format_date(report['last_submission_date'])[0])

    p = doc.add_paragraph("", style="List Number")
    p.add_run("Name of Praesidium: ").bold = True 
    p.add_run(praesidium['name'])

    p = doc.add_paragraph("", style="List Number")
    p.add_run("Address: ").bold = True 
    p.add_run(praesidium['parish'])

    p = doc.add_paragraph("", style="List Number")
    p.add_run("Report No.: ").bold = True 
    p.add_run(str(report['report_number']))
    p.add_run("\t\t\t\t\tDate of Inauguration: ").bold = True
    p.add_run(format_date(praesidium['inaug_date'])[0])

    p = doc.add_paragraph("", style="List Number")
    p.add_run("Day/Time/Venue of Meeting: ").bold = True 
    p.add_run(f"{praesidium['address']}, {praesidium['meeting_time']}")

    p = doc.add_paragraph("", style="List Number")
    p.add_run("Period of Report ").bold = True 
    p.add_run(f"{report['report_period']}")

    p = doc.add_paragraph("", style="List Number")
    p.add_run("Date, Names and Praesidium of Last Curia Visitors: ").bold = True 
    p.add_run(f"{report['last_curia_visitors']} on {format_date(report['last_curia_visit_date'])[0]}")

    pres_term = get_term(praesidium["pres_app_date"], report["submission_date"])
    vp_term = get_term(praesidium['vp_app_date'], report['submission_date']) 
    sec_term = get_term(praesidium['sec_app_date'], report['submission_date'])
    tres_term = get_term(praesidium['tres_app_date'], report['submission_date'])
    officers = [
        ("Office", "Name", "Appointment Date", "Term"), 
        ("Spiritual Director", praesidium['spiritual_director'].title(), format_date(praesidium['spiritual_director_app_date'])[0], ""), 
        ("President", praesidium['president'].title(), format_date(praesidium['pres_app_date'])[0], f'{pres_term}'), 
        ('Vice President', praesidium['vice_president'].title(), format_date(praesidium['vp_app_date'])[0], f'{vp_term}'), 
        ('Secretary', praesidium['secretary'].title(), format_date(praesidium['sec_app_date'])[0], f'{sec_term}'),
        ('Treasurer', praesidium['treasurer'].title(), format_date(praesidium['tres_app_date'])[0], f'{tres_term}')
    ]
    
    doc.add_paragraph()
    p = doc.add_paragraph("", style="List Number")
    p.add_run("Officers Details").bold = True 
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    # Adjust the height of each row
    for i, row in enumerate(table.rows): 
        row.height = Pt(15)
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(2.5)
        row.cells[3].width = Inches(.7)
        for cell, item in zip(row.cells, officers[i]): 
            bold = True if i == 0 else False
            set_cell_content(cell, str(item), bold)

    attendance_to_curia = [
        ("Officer", "No. of Meetings Attended", "Total Meetings for the Period", "Current Year", "Previous Year"), 
        (
            'President',
            report['officers_curia_attendance']['President'], 
            report['no_curia_meetings_held']['President'], 
            f"{report['officers_curia_attendance']['President']} out of {report['no_curia_meetings_held']['President']}", 
            f"{report['previous_curia_attendance']['President']} out of {report['no_curia_meetings_held_previous']['President']}"  if report['no_curia_meetings_held_previous']['President'] else NON 
        ), 
        (
            'Vice President',
            report['officers_curia_attendance']['Vice President'], 
            report['no_curia_meetings_held']['Vice President'], 
            f"{report['officers_curia_attendance']['Vice President']} out of {report['no_curia_meetings_held']['Vice President']}", 
            f"{report['previous_curia_attendance']['Vice President']} out of {report['no_curia_meetings_held_previous']['Vice President']}" if report['no_curia_meetings_held_previous']['Vice President'] else NON
        ), 
        (
            'Secretary',
            report['officers_curia_attendance']['Secretary'], 
            report['no_curia_meetings_held']['Secretary'], 
            f"{report['officers_curia_attendance']['Secretary']} out of {report['no_curia_meetings_held']['Secretary']}", 
            f"{report['previous_curia_attendance']['Secretary']} out of {report['no_curia_meetings_held_previous']['Secretary']}" if report['no_curia_meetings_held_previous']['Secretary'] else NON
        ), 
        (
            'Treasurer',
            report['officers_curia_attendance']['Treasurer'], 
            report['no_curia_meetings_held']['Treasurer'], 
            f"{report['officers_curia_attendance']['Treasurer']} out of {report['no_curia_meetings_held']['Treasurer']}", 
            f"{report['previous_curia_attendance']['Treasurer']} out of {report['no_curia_meetings_held_previous']['Treasurer']}" if report['no_curia_meetings_held_previous']['Treasurer'] else NON
        )
    ]
    doc.add_paragraph()
    p = doc.add_paragraph("", style="List Number")
    p.add_run("Attendance to Curia Meetings").bold = True 
    table = doc.add_table(rows=5, cols=5)
    table.style = "Table Grid"
    for i, row in enumerate(table.rows): 
        row.height = Pt(15)
        for cell, item in zip(row.cells, attendance_to_curia[i]): 
            bold = True if i == 0 else False
            set_cell_content(cell, str(item), bold)

    attendance_to_praesidium = [
        ("Officer", "No. of Meetings Attended", "Total Meetings for the Period", "Current Year", "Previous Year"), 
        (
            'President',
            report['officers_meeting_attendance']['President'], 
            report['no_praesidium_meetings_held']['President'], 
            f"{report['officers_meeting_attendance']['President']} out of {report['no_praesidium_meetings_held']['President']}", 
            f"{report['previous_meeting_attendance']['President']} out of {report['no_praesidium_meetings_held_previous']['President']}" if report['no_praesidium_meetings_held_previous']['President'] else NON 
        ), 
        (
            'Vice President',
            report['officers_meeting_attendance']['Vice President'], 
            report['no_praesidium_meetings_held']['Vice President'], 
            f"{report['officers_meeting_attendance']['Vice President']} out of {report['no_praesidium_meetings_held']['Vice President']}", 
            f"{report['previous_meeting_attendance']['Vice President']} out of {report['no_praesidium_meetings_held_previous']['Vice President']}" if report['no_praesidium_meetings_held_previous']['Vice President'] else NON
        ), 
        (
            'Secretary',
            report['officers_meeting_attendance']['Secretary'], 
            report['no_praesidium_meetings_held']['Secretary'], 
            f"{report['officers_meeting_attendance']['Secretary']} out of {report['no_praesidium_meetings_held']['Secretary']}", 
            f"{report['previous_meeting_attendance']['Secretary']} out of {report['no_praesidium_meetings_held_previous']['Secretary']}" if report['no_praesidium_meetings_held_previous']['Secretary'] else NON
        ), 
        (
            'Treasurer',
            report['officers_meeting_attendance']['Treasurer'], 
            report['no_praesidium_meetings_held']['Treasurer'], 
            f"{report['officers_meeting_attendance']['Treasurer']} out of {report['no_praesidium_meetings_held']['Treasurer']}", 
            f"{report['previous_meeting_attendance']['Treasurer']} out of {report['no_praesidium_meetings_held_previous']['Treasurer']}" if report['no_praesidium_meetings_held_previous']['Treasurer'] else NON
        )
    ]
    doc.add_paragraph()
    p = doc.add_paragraph("", style="List Number")
    p.add_run("Attendance to Praesidium Meetings").bold = True 
    table = doc.add_table(rows=5, cols=5)
    table.style = "Table Grid"
    for i, row in enumerate(table.rows): 
        row.height = Pt(15)
        for cell, item in zip(row.cells, attendance_to_praesidium[i]): 
            bold = True if i == 0 else False
            set_cell_content(cell, str(item), bold)

    membership_mapping = {
        "affiliated_praesidia": "Affiliated Praesidia (if any)",
        "active_members": "Active Members",
        "probationary_members": "Probationary Members",
        "auxiliary_members":"Auxiliary Members" ,
        "adjutorian_members": "Adjutorian Members",
        "praetorian_members": "Praetorian Members"
    }
    report['membership'].pop('id')
    junior_index = 3
    cols = 4
    if not report['include_intermediate']: 
        junior_index = 2
        cols = 3
        membership = [
                (
                    membership_mapping[mem],            
                    report['membership'][mem][0],
                    report['membership'][mem][2] 
                ) for mem in report['membership']
            ]
    else:
        membership = [
                (
                    membership_mapping[mem],            
                    report['membership'][mem][0],
                    report['membership'][mem][1],
                    report['membership'][mem][2] 
                ) for mem in report['membership']
            ]
    memberships_for_total = [
        'active_members', 'probationary_members', 'auxiliary_members'
    ]
    total_seniors = 0 
    for mem in memberships_for_total: 
        total_seniors += report['membership'][mem][0]
    total_intermediates = 0 
    for mem in memberships_for_total: 
        total_intermediates += report['membership'][mem][1]
    total_juniors = 0 
    for mem in memberships_for_total: 
        total_juniors += report['membership'][mem][2]
    
    if report['include_intermediate']:
        membership.append((
            'Total Members', total_seniors, total_intermediates, total_juniors
        )) # type:ignore
    else: 
        membership.append((
            'Total Members', total_seniors, total_juniors
        )) # type:ignore
    doc.add_page_break()
    doc.add_paragraph('')
    p = doc.add_paragraph("", style="List Number")
    p.add_run("Membership").bold = True 
    num_rows = len(membership) + 1
    table = doc.add_table(rows=num_rows, cols=cols)
    table.style = "Table Grid"
    # header row
    row = table.rows[0]
    row.height = Pt(15)
    row.cells[0].width = Inches(2.5)
    set_cell_content(row.cells[0], 'Category', True)
    set_cell_content(row.cells[1], "Senior", True) 
    if report['include_intermediate']:
        set_cell_content(row.cells[2], "Intermediate", True) 
    set_cell_content(row.cells[junior_index], "Junior", True)
    for i in range(1, len(table.rows)):
        row = table.rows[i]
        row.height = Pt(15)
        row.cells[0].width = Inches(2.5)
        for cell, item in zip(row.cells, membership[i-1]): 
            bold = True if "Total" in str(item) else False
            set_cell_content(cell, str(item), bold)


    doc.add_paragraph('')
    p = doc.add_paragraph("", style="List Number") 
    p.add_run("Meetings and Attendance").bold = True 
    p = doc.add_paragraph("")
    p.add_run("No. of meetings expected to be held: ").bold = True 
    p.add_run(f"{report['no_meetings_expected']}")
    p = doc.add_paragraph("")
    p.add_run("No. of meetings held: ").bold = True 
    p.add_run(f"{report['no_meetings_held']}")
    p = doc.add_paragraph("")
    p.add_run("Average attendance per meeting: ").bold = True 
    p.add_run(f"{report['avg_attendance']}")
    p = doc.add_paragraph("")
    p.add_run("Reason for poor attendance: ").bold = True 
    p.add_run(f"{report['poor_attendance_reason']}")

    # Works
    active_works = list(filter(lambda item: item['active'], report['work_summary']))
    other_works = list(filter(lambda item: not item['active'], report['work_summary']))

    doc.add_paragraph('')
    p = doc.add_paragraph("", style="List Number") 
    p.add_run("Works Undertaken").bold = True 
    p1 = doc.add_paragraph()
    p1.add_run('Active: ').bold = True
    p1.add_run(', '.join([work['type'] for work in active_works]))

    p2 = doc.add_paragraph()
    p2.add_run('Others: ').bold = True
    p2.add_run(', '.join([work['type'] for work in other_works]))


    # Active works 
    no_of_rows = ceil(len(active_works) / 2)
    table = doc.add_table(rows=no_of_rows, cols=2)
    table.style = "Table Grid"

    # run_2.font.size = Pt(10)  # Adjust font size

    work_count = 65 # Letter 'A'
    for i, work in enumerate(active_works):
        col = i % 2; row = i // 2
        cell = table.cell(row, col)
        p = cell.paragraphs[0]
        run1 = p.add_run(f'\n{chr(work_count)}. {work["type"]}\n')
        run1.bold = True

        key_val = list(work['details'].items())

        # Incude total or average
        calc_total_or_avg = report['work_total_and_average'].get(work['type'], None)
        total = 0; average = 0
        calc_total = False
        calc_average = False
        if calc_total_or_avg: 
            total = sum([val for (key, val) in key_val if not ('home' in key)])
            if calc_total_or_avg['total']: 
                calc_total = True
            if calc_total_or_avg['average']: 
                calc_average = True
                average = ceil(total / work['no_done'])
        
        # cell_text = '' # f'{chr(work_count)}. {work["type"]}\n\n'
        cell_text = f"Done: {work['no_done']}, Assigned: {work['no_assigned']}\n"
        if calc_total: cell_text += f"Total: {total} "
        if calc_average: cell_text += f", Average: {average}"
        if calc_total_or_avg: cell_text += '\n'
        # Include other metrics
        for a, b in key_val: 
            cell_text += str(a) + ": " + str(b) + '\n'
        
        run2 = p.add_run(cell_text)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        work_count += 1

    # Other works 
    doc.add_page_break()
    # doc.add_paragraph('')
    p = doc.add_paragraph("", style="List Number") 
    p.add_run("Other Works").bold = True 
    
    no_of_rows = len(other_works)
    table = doc.add_table(rows=no_of_rows+1, cols=3)
    table.style = "Table Grid"
    for row in table.rows: 
        row.height = Pt(15)    
    set_cell_content(table.cell(0,0), "Work", True)
    set_cell_content(table.cell(0,1), "No. assigned", True)
    set_cell_content(table.cell(0,2), "No. done", True)

    work_count = 65
    for i, work in enumerate(other_works):
        # col = i % 2; 
        row = i + 1
        for i, key in enumerate(['type', 'no_assigned', 'no_done']):
            cell = table.cell(row, i)
            set_cell_content(cell, str(work[key]))

    # Achievements
    achievement_mapping = {
        'Category': 'Category',
        'no_baptized': 'Baptized', 
        'no_confirmed': "Confirmed", 
        'no_converted': "Converted to Christianity", 
        'no_first_communicants': 'First communicants', 
        'no_married': "Marriages solemnized", 
        'no_recruited': "Recruited into the Legion", 
        'no_vocations': "Called to vocation", 
        'no_promised': "Legion promises taken"
    }
    achievements = {
        'Category': ['Current year', 'Previous year']
    }
    achievements.update(report['achievements'])
    achievements.pop('id')
    other_key_pairs = list(achievements['others'].items())
    other_key_pairs = other_key_pairs[0] if other_key_pairs else []
    achievements.pop('others')
    if other_key_pairs:
        achievement_mapping.update({'others': other_key_pairs[0]})
        achievements.update({'others': other_key_pairs[1]})
    num_rows = len(achievements)
    
    # print('other_key_paris', other_key_pairs, achievements, achievement_mapping)

    doc.add_paragraph('')
    p = doc.add_paragraph("", style="List Number") 
    p.add_run("Achievements Recorded").bold = True 
    table = doc.add_table(rows=num_rows, cols=3)
    table.style = "Table Grid"
    for i, key in enumerate(achievements.keys()):
        row = table.rows[i]
        row.height = Pt(15)
        bold = True if i == 0 else False
        set_cell_content(row.cells[0], achievement_mapping[key], bold)
        set_cell_content(row.cells[1], str(achievements[key][0]), bold)
        set_cell_content(row.cells[2], str(achievements[key][1]), bold)


    # Plans for extension
    doc.add_paragraph('')
    p = doc.add_paragraph("", style="List Number") 
    p.add_run("Plans for Extension Work").bold = True 
    doc.add_paragraph(report['extension_plans'])

    # Functions 
    doc.add_paragraph('')
    p = doc.add_paragraph('', style="List Number")
    p.add_run('Legion Functions with Attendance').bold = True
    num_rows = len(report['fxn_attendances'])+1
    table = doc.add_table(rows=num_rows, cols=5)
    table.style = "Table Grid"
    # header row
    row = table.rows[0]
    row.height = Pt(15)
    row.cells[0].width = Inches(.6)
    set_cell_content(row.cells[0], 'S/N', True)
    set_cell_content(row.cells[1], 'Function', True)
    set_cell_content(row.cells[2], "Date", True) 
    set_cell_content(row.cells[3], "Current year", True)
    set_cell_content(row.cells[4], "Previous year", True)

    special_fxns = ["Patrician Meetings", "October Devotion", "May Devotion"]

    for i, fxn in enumerate(report['fxn_attendances']):
        row = table.rows[i+1]
        row.height = Pt(15)
        row.cells[0].width = Inches(.5)
        row.cells[1].width = Inches(2.1)
        row.cells[2].width = Inches(1.8)
        set_cell_content(row.cells[0], chr(97+i))
        set_cell_content(row.cells[1], fxn['name'])
        if fxn['name'] not in special_fxns:
            set_cell_content(row.cells[2], format_date(fxn['date'])[0])
        elif fxn['name'] == "Patrician Meetings": 
            month_range = format_date(report['patricians_start'])[0] + ' to ' + format_date(report['patricians_end'])[0]
            set_cell_content(row.cells[2], month_range)
        elif fxn['name'] == "October Devotion": 
            year = format_date(fxn['date'])[3]
            set_cell_content(row.cells[2], f"1 Oct to 31 Oct, {year}")
        elif fxn['name'] == "May Devotion": 
            year = format_date(fxn['date'])[3]
            set_cell_content(row.cells[2], f"1 May to 31 May, {year}")

        set_cell_content(row.cells[3], str(fxn['current_year_attendance']))
        set_cell_content(row.cells[4], str(fxn['previous_year_attendance']))


    # Finance
    finances = report['financial_summary']
    total_income = finances['acf'][0] + sum(finances['sbc'])
    income_dict = {
        "Income": NULL, 
        "Amount brought forward from last report": finances['acf'][0], 
        "SBC for the period": sum(finances['sbc']), 
        "Total Income": total_income
    }
    others = finances['expenses']['others']
    remittance = sum(finances['expenses']['remittance'])
    bouquet = sum(finances['expenses']['bouquet'])
    stationery = sum(finances['expenses']['stationery'])
    altar = sum(finances['expenses']['altar'])
    extension = sum(finances['expenses']['extension'])
    production = finances['report_production']
    others = sum([val for arr in others for obj in arr for (_, val) in obj.items()])
    total_expenses = remittance + bouquet + stationery + altar + extension + production + others
    expenses_dict = {
        "Expenditure": NULL,
        "Monthly Remittance": remittance, 
        "Spiritual Bouquet": bouquet, 
        "Stationery": stationery, 
        "Altar": altar, 
        "Extension": extension, 
        "Production of Report": production, 
        "Others": others,
        "Total Expenses": total_expenses, 
        "Surplus Funds to Curia": total_income - total_expenses, 
        "Balance at Hand": finances['balance_at_hand']
    }
    doc.add_page_break()
    doc.add_paragraph('')
    p = doc.add_paragraph('', style="List Number")
    p.add_run('Finance').bold = True 
    num_rows = len(income_dict)+len(expenses_dict)+1
    income_shift = len(income_dict) + 1

    table = doc.add_table(rows=num_rows, cols=3)
    table.style = "Table Grid"
    row = table.rows[0]     
    row.height = Pt(15)
    cell = row.cells[0]
    set_cell_content(cell, 'Description', True)
    set_cell_content(row.cells[1], currency, True)
    set_cell_content(row.cells[2], currency, True)

    for i in [j+1 for j in range(num_rows-1)]:   
        row = table.rows[i]    
        row.height = Pt(15)
        if i < income_shift: 
            # handle income 
            keys, vals = list(income_dict.keys()), list(income_dict.values())
            key, val = keys[i-1], vals[i-1]
            bold = True if ('Desc' in key or 'Income' in key) else False
            set_cell_content(row.cells[0], str(key), bold)
            set_cell_content(row.cells[2], str(val))
        elif i < (num_rows-3): 
            # handle expenditure 
            keys, vals = list(expenses_dict.keys()), list(expenses_dict.values())
            key, val = keys[i-income_shift], vals[i-income_shift]
            bold = True if ('Expenditure' in key or 'Total' in key or 'Surplus' in key) else False
            set_cell_content(row.cells[0], str(key), bold)
            set_cell_content(row.cells[1], str(val))
        else: 
            keys, vals = list(expenses_dict.keys()), list(expenses_dict.values())
            key, val = keys[i-income_shift], vals[i-income_shift]
            bold = True 
            set_cell_content(row.cells[0], str(key), bold)
            set_cell_content(row.cells[2], str(val))


    doc.add_paragraph('')
    p = doc.add_paragraph('', style="List Number")
    p.add_run('Was the finance audited? ').bold = True 
    p.add_run('Yes' if report['audited'] else 'No')

    p = doc.add_paragraph('', style="List Number")
    p.add_run('Problems: ').bold = True 
    p.add_run(report['problems'])

    p = doc.add_paragraph('', style="List Number")
    p.add_run('Was the report read and accepted by members before submission? ').bold = True 
    p.add_run('Yes' if report['read_and_accepted'] else 'No')

    p = doc.add_paragraph('', style="List Number")
    p.add_run('Any other comment/remark: ').bold = True 
    p.add_run(report['remarks'])

    p = doc.add_paragraph('', style="List Number")
    p.add_run('Conclusion: ').bold = True 
    p.add_run(report['conclusion'])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p1 = p.add_run('.'*65 + '\t\t\t' + '.'*65)

    p = doc.add_paragraph()
    p1 = p.add_run(praesidium['president'] + '\t'*5 + praesidium['secretary'])
    p1.bold = True

    p = doc.add_paragraph()
    p1 = p.add_run('\tPresident' + '\t'*6 + 'Secretary')
    p1.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p1 = p.add_run("\nSpiritual Director's Comment:")
    p1.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('.'*516).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph('.'*60)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph()
    p1 = p.add_run(praesidium['spiritual_director'] + '\nSpiritual Director')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p1.bold = True

    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("Auditor's Report")
    r.bold = True 
    r.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = doc.add_table(rows=num_rows, cols=7)
    table.style = "Table Grid"
    row = table.rows[0]     
    row.height = Pt(15)
    set_cell_content(row.cells[2], 'Income', True)
    set_cell_content(row.cells[4], 'Expenditure', True)
    row = table.rows[1]     
    row.height = Pt(15)
    set_cell_content(row.cells[0], 'S/N', True)
    row.cells[0].width = Inches(.6)
    set_cell_content(row.cells[1], 'Month', True)
    set_cell_content(row.cells[2], f'BF ({currency})', True)
    set_cell_content(row.cells[3], f'SBC ({currency})', True)
    set_cell_content(row.cells[4], f'Curia ({currency})', True)
    set_cell_content(row.cells[5], f'Praesidium ({currency})', True)
    set_cell_content(row.cells[6], f'Balance ({currency})', True)



    try: 
        doc.save('test_rep4.docx')
    except PermissionError: 
        print("A file in this location already has that name")



generate_report_docx(curia, praesidium, report)